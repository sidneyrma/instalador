# -*- coding: utf-8 -*-
"""
O Ouro das Palavras — Joseph Murphy
Compilação, limpeza editorial e diagramação em livro (.docx) conforme ABNT NBR 6029.
Gera também um PDF-espelho para calcular a paginação do sumário.
"""
import re, os, json, sys
from pathlib import Path

HERE = Path(__file__).parent
FONTE = HERE / "fonte_original.txt"
FONTES_DIR = HERE / "fontes"
OUT_DOCX = HERE / "O_Ouro_das_Palavras.docx"
OUT_PDF = HERE / "O_Ouro_das_Palavras_previa.pdf"

# ---------------------------------------------------------------------------
# 1. LIMPEZA EDITORIAL
# ---------------------------------------------------------------------------
def limpar(t: str) -> str:
    # marcas d'água do transcritor (parágrafos inteiros)
    t = re.sub(r'\(Este arquivo tem mais de 30 minutos\.?[^\n]*\)', '', t)
    t = re.sub(r'\(Transcrito por TurboScribe\.?[^\n]*\)', '', t)
    # chamada comercial do canal (YouTube)
    t = t.replace(
        "Observamos que mais de 85% dos que visitam nosso canal não estão inscritos. "
        "Você pode apoiar este projeto pressionando o botão de inscrição e curtindo o vídeo. "
        "Obrigado por nos inspirar a continuar criando mais e melhor conteúdo para você.", "")
    t = t.replace("Você está pronto para transformar sua vida? ", "")
    # reparos de transcrição (transcritos fielmente do original, com ajustes pontuais)
    t = t.replace("Doutor. Morphe.", "Doutor Murphy,")
    t = t.replace("Salmo 104, 3.", "Salmo 104, 33.")          # citação correta
    t = t.replace("Mateus 6,20", "Mateus 18, 20")             # citação correta
    t = t.replace("Provérbios 10, 2", "Provérbios 10, 22")    # citação correta
    t = t.replace("Quando no Gênesis lemos Assim deve ser o seu, o uso da palavra.",
                  "Quando no Gênesis lemos: assim deve ser o seu uso da palavra.")
    t = t.replace("Cole em suas condições.", "Colha em suas condições.")
    t = t.replace("muitos se resistem", "muitos resistem")
    t = t.replace("Isso amplificará o efeito de sua gestão.",
                  "Isso amplificará o efeito de sua afirmação.")
    t = t.replace("Este ato imaginário, Imaginalate,",
                  "Este ato imaginário, o ato imaginal,")
    t = t.replace("erva da ninha", "erva daninha")
    t = t.replace("Esse silêncio é como a pousada na agricultura.",
                  "Esse silêncio é como o pousio na agricultura.")
    # nomes próprios corrigidos (transcrição de áudio)
    t = t.replace("Antoine Dubois", "Émile Coué")
    t = t.replace("Dmitry Volkov", "Ivan Pavlov")
    t = t.replace("Neville Boddard", "Neville Goddard")
    t = t.replace("Charles Hanel", "Charles Haanel")
    t = t.replace("Carl Dunn", "Carl Jung")
    t = t.replace("método Kowe", "método Coué")
    t = t.replace("cantos de Taiki", "cantos de Taizé")
    t = t.replace("O que todos esses os casos demonstram a mesma coisa.",
                  "O que todos esses casos demonstram é a mesma coisa.")
    t = t.replace("Passam silêncio.", "Passem em silêncio.")
    t = t.replace("a caridade vibracional compartilhada", "a qualidade vibracional compartilhada")
    t = re.sub(r'\n{3,}', '\n\n', t)
    return t.strip()

# ---------------------------------------------------------------------------
# 2. ESTRUTURA — capítulos
# ---------------------------------------------------------------------------
CAPITULOS = [
    # (número, marcador de início, início do corpo, título)
    (1,  "Capítulo 1 – O Grito Inicial", "No princípio era a palavra", "O Grito Inicial"),
    (2,  "Capítulo II A Vibração da Palavra Pensada", "Toda palavra é vibração", "A Vibração da Palavra Pensada"),
    (3,  "Capítulo terceiro, o laboratório interior", "Dentro de você existe uma oficina oculta", "O Laboratório Interior"),
    (4,  "Capítulo 4. Frases de poder para a saúde", "Cada célula do seu corpo obedece", "Frases de Poder para a Saúde"),
    (5,  "A abundância não é um acaso", "A abundância não é um acaso", "Frases de Poder para a Abundância"),
    (6,  "Capítulo Sexto Frases de Poder para Relacionamentos", "Todo vínculo humano", "Frases de Poder para Relacionamentos"),
    (7,  "Capítulo 7. Quando o decreto parece falhar", "Aquele que começou a usar o poder do decreto", "Quando o Decreto Parece Falhar"),
    (8,  "Existe uma força que transcende", "Existe uma força que transcende", "O Decreto Coletivo"),
    (9,  "Capítulo 9 Transmutar crises em plataformas", "Não há crise que não possa se tornar um alicerce", "Transmutar Crises em Plataformas"),
    (10, "Capítulo 10 Um Ritual Diário do Verbo Dourado", "Todo poder espiritual, para se tornar transformação concreta", "Um Ritual Diário do Verbo Dourado"),
]

INTRO_INICIO = "Vivemos em um mundo que reverencia o que é visível."
INTRO_FIM = ("Comece agora o Ouro das Palavras baseado nos ensinamentos de Joseph Murphy, "
             "editado e publicado por Coleção do Despertar. Todos os direitos reservados.")

def segmentar(t: str):
    assert INTRO_INICIO in t, "início da introdução não encontrado"
    assert INTRO_FIM in t, "final da introdução não encontrado"
    i0 = t.index(INTRO_INICIO)
    i1 = t.index(INTRO_FIM) + len(INTRO_FIM)
    intro = t[i0:i1].replace(INTRO_FIM, "Comece agora.")
    rest = t[i1:]

    cap = []
    for idx, (n, marc, corpo_inicio, titulo) in enumerate(CAPITULOS):
        j = rest.find(marc)
        assert j >= 0, f"marcador do capítulo {n} não encontrado: {marc[:40]}"
        rest = rest[j:]
        k = rest.find(corpo_inicio)
        assert k >= 0, f"corpo do capítulo {n} não encontrado"
        # fim do capítulo = início do próximo marcador
        prox = CAPITULOS[idx + 1][1] if idx + 1 < len(CAPITULOS) else None
        if prox:
            fim = rest.find(prox)
            assert fim >= 0, f"fim do capítulo {n} não encontrado"
            corpo = rest[k:fim]
        else:
            corpo = rest[k:]
        cap.append({"num": n, "titulo": titulo, "corpo": corpo.strip()})
        rest = rest[fim:] if prox else ""
    return intro, cap

# ---------------------------------------------------------------------------
# 3. AJUSTES POR CAPÍTULO (títulos internos e destaques)
# ---------------------------------------------------------------------------
def ajustar(cap: dict) -> list:
    """Devolve lista de blocos: (tipo, texto)  tipo ∈ {h1num, h1, h2, p}"""
    n, corpo = cap["num"], cap["corpo"]
    blocos = [("h1num", f"CAPÍTULO {n}"), ("h1", cap["titulo"])]

    if n == 2:  # resumo com numeração ordinal
        corpo = corpo.replace("Em resumo, primeiro,", "Em resumo, **primeiro**,")
        for kw in ("Segundo,", "Terceiro,", "Quarto,", "Quinto,"):
            corpo = re.sub(r'(?<![\w])(' + kw + r') ', r'**\1** ', corpo)

    if n == 4:  # frases por sistema do corpo
        corpo = corpo.replace(
            "Aqui estão algumas frases de poder para aplicar para o sistema nervoso.",
            "Aqui estão algumas frases de poder, direcionadas a cada área do corpo: "
            "**Sistema nervoso.**")
        rotulos = [
            ("para o sistema nervoso.", "**Sistema nervoso.**"),
            ("Para o sistema digestivo,", "**Sistema digestivo.**"),
            ("Para o sistema respiratório,", "**Sistema respiratório.**"),
            ("Para o sistema imunológico,", "**Sistema imunológico.**"),
            ("Para a pele e tecidos,", "**Pele e tecidos.**"),
            ("Para os olhos,", "**Olhos.**"),
            ("Para os ossos e articulações,", "**Ossos e articulações.**"),
            ("Para o sistema endócrino,", "**Sistema endócrino.**"),
            ("Para os órgãos sexuais e reprodutivos,", "**Órgãos sexuais e reprodutivos.**"),
        ]
        for a, b in rotulos:
            corpo = corpo.replace(a, b)

    if n == 5:  # frases para a abundância
        corpo = corpo.replace("Use-as com reverência para ativar o fluxo geral.",
                              "Use-as com reverência. **Ativar o fluxo geral.**")
        trocas = [
            ("Tudo o que toco prospera para transformar dívida ou carência.",
             "Tudo o que toco prospera. **Transformar dívida ou carência.**"),
            ("Eu recebo com alegria para me abrir a oportunidades.",
             "Eu recebo com alegria. **Abrir-se a oportunidades.**"),
            ("As oportunidades me encontram para me afirmar como canal.",
             "As oportunidades me encontram. **Afirmar-se como canal.**"),
            ("Eu circulo o bem e o recebo multiplicado para agradecer antecipadamente.",
             "Eu circulo o bem e o recebo multiplicado. **Agradecer antecipadamente.**"),
        ]
        for a, b in trocas:
            corpo = corpo.replace(a, b)

    if n == 7:  # método de neutralização e reinício
        corpo = corpo.replace("Reconhecimento. O primeiro passo", "**Reconhecimento.** O primeiro passo")
        corpo = corpo.replace("Silêncio. Uma vez reconhecido", "**Silêncio.** Uma vez reconhecido")
        corpo = corpo.replace("Reprogramação. Após o silêncio", "**Reprogramação.** Após o silêncio")

    if n == 9:  # categorias de frases
        trocas = [
            ("respire-as, para perdas materiais.", "respire-as. **Perdas materiais.**"),
            ("Para doenças ou diagnósticos adversos,", "**Doenças ou diagnósticos adversos.**"),
            ("Para rupturas ou separações,", "**Rupturas ou separações.**"),
            ("Para rejeições ou fracassos profissionais,", "**Rejeições ou fracassos profissionais.**"),
            ("Para crises existenciais ou momentos de vazio,", "**Crises existenciais ou momentos de vazio.**"),
        ]
        for a, b in trocas:
            corpo = corpo.replace(a, b)

    if n == 10:  # subseções do ritual diário
        Q = chr(10) + chr(10)          # quebra de parágrafo real
        def h2(titulo, texto):
            return "[[H2]]" + titulo + "[[/H2]]" + Q + texto
        trocas = [
            ("Cada um cumpre uma função diferente e cada um é indispensável. E ao despertar a semente, o primeiro pensamento do dia",
             "Cada um cumpre uma função diferente e cada um é indispensável." + Q + h2("1. Ao despertar: a semente", "O primeiro pensamento do dia")),
            ("Levante-se. As duas ao meio-dia do reforço. No meio do dia",
             "Levante-se." + Q + h2("2. Ao meio-dia: o reforço", "No meio do dia")),
            ("Com ela, o novo padrão é reforçado. 3. Antes de dormir, a integração.",
             "Com ela, o novo padrão é reforçado." + Q + h2("3. Antes de dormir: a integração", "")),
            ("Entregue-se ao sono com confiança. O subconsciente obedecerá. O ciclo de 30 dias.",
             "Entregue-se ao sono com confiança. O subconsciente obedecerá." + Q + h2("O ciclo de 30 dias", "")),
        ]
        for a, b in trocas:
            corpo = corpo.replace(a, b)

    # capitaliza a palavra que segue um rótulo em negrito (ex.: "**Sistema respiratório.** o sopro")
    corpo = re.sub(r'(\*\*[^*]+\*\*\s+)([a-zçáéíóúâêôãõà])',
                   lambda m: m.group(1) + m.group(2).upper(), corpo)

    # quebra em parágrafos
    for par in corpo.split("\n\n"):
        par = par.strip()
        if not par:
            continue
        if par.startswith("[[H2]]") and par.endswith("[[/H2]]"):
            blocos.append(("h2", par[6:-7].strip()))
        else:
            blocos.append(("p", par))
    return blocos

# ---------------------------------------------------------------------------
# 4. CONTEÚDO FINAL (blocos)
# ---------------------------------------------------------------------------
def montar_conteudo(paginas=None):
    paginas = paginas or {}
    t = limpar(FONTE.read_text(encoding="utf-8"))
    intro, caps = segmentar(t)
    # lista de seções para o sumário
    secao_sumar = [("Introdução", "introducao")] + [(f"Capítulo {c['num']} — {c['titulo']}", f"cap{c['num']}") for c in caps]
    conteudo = {
        "introducao": [("h1", "Introdução")] + [("p", p) for p in intro.split("\n\n") if p.strip()],
        "capitulos": [ajustar(c) for c in caps],
        "sumario": secao_sumar,
        "paginas": paginas,
    }
    return conteudo

# ---------------------------------------------------------------------------
# 5. GERAÇÃO .DOCX (python-docx)
# ---------------------------------------------------------------------------
def gerar_docx(conteudo, caminho):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    TNR = "Times New Roman"

    def estilo_base(st):
        st.font.name = TNR
        st.font.size = Pt(12)
        rpr = st.element.get_or_add_rPr()
        rf = rpr.find(qn('w:rFonts'))
        if rf is None:
            rf = OxmlElement('w:rFonts'); rpr.append(rf)
        for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
            rf.set(qn(a), TNR)
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Cm(1.25)

    doc = Document()
    estilo_base(doc.styles['Normal'])

    # propriedades do documento
    cp = doc.core_properties
    cp.title = "O Ouro das Palavras"
    cp.author = "Joseph Murphy"
    cp.subject = "O poder criador da palavra"
    cp.keywords = "Joseph Murphy; palavras; decretos; subconsciente; espiritualidade"
    cp.comments = "Baseado nos ensinamentos de Joseph Murphy. Editado e publicado por Coleção do Despertar."

    def config_sec(sec):
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.top_margin = Cm(3.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)
        sec.header_distance = Cm(1.5)
        sec.footer_distance = Cm(1.5)

    secA = doc.sections[0]
    config_sec(secA)
    secA.different_first_page_header_footer = True  # capa sem cabeçalho

    def add_par(texto=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False,
                italic=False, indent=None, before=0, after=0, line=None,
                page_break=False, keep_next=False):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = align
        if indent is not None:
            pf.first_line_indent = indent
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        if line:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(line)
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if page_break:
            pf.page_break_before = True
        if keep_next:
            pf.keep_with_next = True
        if texto:
            r = p.add_run(texto)
            r.font.name = TNR
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.italic = italic
        return p

    def add_rich(texto, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, indent=Cm(1.25),
                 before=0, after=0, keep_next=False):
        """texto com marcação **negrito** inline"""
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = align
        if indent is not None:
            pf.first_line_indent = indent
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        if keep_next:
            pf.keep_with_next = True
        partes = re.split(r'(\*\*.*?\*\*)', texto)
        for parte in partes:
            if not parte:
                continue
            negrito = parte.startswith('**') and parte.endswith('**')
            txt = parte[2:-2] if negrito else parte
            r = p.add_run(txt)
            r.font.name = TNR
            r.font.size = Pt(size)
            r.font.bold = negrito
        return p

    def add_sumario_linha(titulo, pagina):
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(titulo); r.font.name = TNR; r.font.size = Pt(12)
        r2 = p.add_run("\t"); r2.font.name = TNR; r2.font.size = Pt(12)
        r3 = p.add_run(str(pagina)); r3.font.name = TNR; r3.font.size = Pt(12)
        return p

    def vazios(n):
        for _ in range(n):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(18)

    # ===== CAPA =====
    add_par("COLEÇÃO DO DESPERTAR", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True, after=0)
    vazios(5)
    add_par("O OURO DAS PALAVRAS", align=WD_ALIGN_PARAGRAPH.CENTER, size=30, bold=True, after=12)
    add_par("O poder criador da palavra", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, italic=True, after=6)
    vazios(4)
    add_par("JOSEPH MURPHY", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True)
    vazios(6)
    add_par("COLEÇÃO DO DESPERTAR", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=6)
    add_par("2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    # ===== FOLHA DE ROSTO =====
    add_par("JOSEPH MURPHY", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, page_break=True, after=0)
    vazios(6)
    add_par("O Ouro das Palavras", align=WD_ALIGN_PARAGRAPH.CENTER, size=18, bold=True, after=10)
    add_par("O poder criador da palavra", align=WD_ALIGN_PARAGRAPH.CENTER, size=13, italic=True)
    vazios(9)
    add_par("Coleção do Despertar", align=WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    add_par("2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

    # ===== CRÉDITOS =====
    add_par("© Coleção do Despertar, 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=10, page_break=True, after=0)
    vazios(3)
    add_par("O Ouro das Palavras", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, after=8)
    add_par("Baseado nos ensinamentos de Joseph Murphy.", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=6)
    add_par("Editado e publicado por Coleção do Despertar.", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=6)
    add_par("Todos os direitos reservados.", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=6)
    add_par("Brasil · 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, after=0)
    vazios(3)
    add_par("Diagramação e formatação segundo a ABNT NBR 6029 —", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, after=2)
    add_par("Apresentação de livros e folhetos.", align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

    # ===== SUMÁRIO =====
    add_par("SUMÁRIO", align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, page_break=True, after=24, keep_next=True)
    for titulo, chave in conteudo["sumario"]:
        add_sumario_linha(titulo, conteudo["paginas"].get(chave, 0))

    # ===== CORPO (seção B: numeração inicia) =====
    secB = doc.add_section(WD_SECTION.NEW_PAGE)
    config_sec(secB)
    secB.different_first_page_header_footer = False
    # reiniciar numeração em 1
    sectPr = secB._sectPr
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    pg = OxmlElement('w:pgNumType'); pg.set(qn('w:start'), '1')
    sectPr.append(pg)
    # cabeçalho com número de página à direita
    hdr = secB.header
    hdr.is_linked_to_previous = False
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    r.font.name = TNR; r.font.size = Pt(12)

    # introdução
    for tipo, txt in conteudo["introducao"]:
        if tipo == "h1":
            add_par(txt, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=24, keep_next=True)
        else:
            add_rich(txt)

    # capítulos
    for blocos in conteudo["capitulos"]:
        for tipo, txt in blocos:
            if tipo == "h1num":
                add_par(txt, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, bold=True,
                        after=6, page_break=True, keep_next=True)
            elif tipo == "h1":
                add_par(txt, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, bold=True, after=24, keep_next=True)
            elif tipo == "h2":
                add_par(txt, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, bold=True,
                        before=12, after=6, keep_next=True)
            else:
                add_rich(txt)

    doc.save(caminho)
    print("DOCX gerado:", caminho)

# ---------------------------------------------------------------------------
# 6. GERAÇÃO PDF-ESPELHO (reportlab) — para paginação
# ---------------------------------------------------------------------------
def gerar_pdf(conteudo, caminho, numero_paginas=False, inicio_corpo=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_LEFT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, PageBreak, KeepTogether)
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    F = str(FONTES_DIR)
    pdfmetrics.registerFont(TTFont('TNR', os.path.join(F, 'Tinos-Regular.ttf')))
    pdfmetrics.registerFont(TTFont('TNR-Bold', os.path.join(F, 'Tinos-Bold.ttf')))
    pdfmetrics.registerFont(TTFont('TNR-Italic', os.path.join(F, 'Tinos-Italic.ttf')))
    pdfmetrics.registerFont(TTFont('TNR-BoldItalic', os.path.join(F, 'Tinos-BoldItalic.ttf')))
    pdfmetrics.registerFontFamily('TNR', normal='TNR', bold='TNR-Bold',
                                  italic='TNR-Italic', boldItalic='TNR-BoldItalic')

    LW, LH = A4  # 595.27 x 841.89 pt
    M_ESQ, M_DIR, M_SUP, M_INF = 3*cm, 2*cm, 3*cm, 2*cm

    corpo = ParagraphStyle('corpo', fontName='TNR', fontSize=12, leading=20.7,
                           alignment=TA_JUSTIFY, firstLineIndent=1.25*cm,
                           spaceBefore=0, spaceAfter=0)
    h1num = ParagraphStyle('h1num', parent=corpo, alignment=TA_CENTER, firstLineIndent=0,
                           fontSize=12, leading=18, spaceAfter=6, keepWithNext=True)
    h1 = ParagraphStyle('h1', parent=corpo, alignment=TA_CENTER, firstLineIndent=0,
                        fontSize=14, leading=21, spaceAfter=24, keepWithNext=True)
    h2 = ParagraphStyle('h2', parent=corpo, alignment=TA_LEFT, firstLineIndent=0,
                        fontSize=12, leading=18, spaceBefore=12, spaceAfter=6, keepWithNext=True)
    cent = ParagraphStyle('cent', parent=corpo, alignment=TA_CENTER, firstLineIndent=0)

    def cpar(txt, style, size=None, bold=False, italic=False):
        st = ParagraphStyle('x', parent=style)
        if size: st.fontSize = size; st.leading = size * 1.2
        st.fontName = ('TNR-Bold' if bold else 'TNR-Italic' if italic else 'TNR')
        return Paragraph(txt, st)

    def vazio(pts):
        return Spacer(1, pts)

    def rodape_num(canvas, doc):
        if inicio_corpo is None or numero_paginas is False:
            return
        pag = canvas.getPageNumber()
        if pag >= inicio_corpo:
            canvas.saveState()
            canvas.setFont('TNR', 12)
            canvas.drawRightString(LW - M_DIR, LH - 1.5*cm, str(pag - inicio_corpo + 1))
            canvas.restoreState()

    doc = SimpleDocTemplate(str(caminho), pagesize=A4,
                            leftMargin=M_ESQ, rightMargin=M_DIR,
                            topMargin=M_SUP, bottomMargin=M_INF,
                            title="O Ouro das Palavras", author="Joseph Murphy")
    story = []

    # capa
    story += [vazio(4.2*cm),
              cpar("COLEÇÃO DO DESPERTAR", cent, 12, bold=True),
              vazio(2.4*cm),
              cpar("O OURO DAS PALAVRAS", cent, 30, bold=True),
              cpar("O poder criador da palavra", cent, 16, italic=True),
              vazio(1.8*cm),
              cpar("JOSEPH MURPHY", cent, 18, bold=True),
              vazio(3.2*cm),
              cpar("COLEÇÃO DO DESPERTAR", cent, 11),
              cpar("2026", cent, 11),
              PageBreak()]
    # folha de rosto
    story += [vazio(1*cm),
              cpar("JOSEPH MURPHY", cent, 14, bold=True),
              vazio(3*cm),
              cpar("O Ouro das Palavras", cent, 18, bold=True),
              cpar("O poder criador da palavra", cent, 13, italic=True),
              vazio(4.2*cm),
              cpar("Coleção do Despertar", cent, 12),
              cpar("2026", cent, 12),
              PageBreak()]
    # créditos
    story += [vazio(1.6*cm),
              cpar("© Coleção do Despertar, 2026", cent, 10),
              vazio(1*cm),
              cpar("O Ouro das Palavras", cent, 11, bold=True),
              cpar("Baseado nos ensinamentos de Joseph Murphy.", cent, 11),
              cpar("Editado e publicado por Coleção do Despertar.", cent, 11),
              cpar("Todos os direitos reservados.", cent, 11),
              cpar("Brasil · 2026", cent, 11),
              vazio(1*cm),
              cpar("Diagramação e formatação segundo a ABNT NBR 6029 —", cent, 9),
              cpar("Apresentação de livros e folhetos.", cent, 9),
              PageBreak()]

    # sumário
    st_sum = ParagraphStyle('sum', parent=corpo, alignment=TA_LEFT, firstLineIndent=0,
                            leftIndent=0)
    try:
        st_sum.tablocs = [(16*cm, 'R', '.')]
        usa_tab = True
    except Exception:
        usa_tab = False
    story.append(cpar("SUMÁRIO", h1, 14, bold=True))
    for titulo, chave in conteudo["sumario"]:
        num = conteudo["paginas"].get(chave, 0)
        if usa_tab:
            story.append(Paragraph(f'{titulo}<tab/>{num}', st_sum))
        else:
            pts = '.' * max(2, 60 - len(titulo))
            story.append(Paragraph(f'{titulo} {pts} {num}', st_sum))
    story.append(PageBreak())

    # corpo
    for tipo, txt in conteudo["introducao"]:
        if tipo == "h1":
            story.append(cpar(txt, h1, 14, bold=True))
        else:
            story.append(Paragraph(html_bold(txt), corpo))
    for blocos in conteudo["capitulos"]:
        story.append(PageBreak())
        for tipo, txt in blocos:
            if tipo == "h1num":
                story.append(cpar(txt, h1num, 12, bold=True))
            elif tipo == "h1":
                story.append(cpar(txt, h1, 14, bold=True))
            elif tipo == "h2":
                story.append(cpar(txt, h2, 12, bold=True))
            else:
                story.append(Paragraph(html_bold(txt), corpo))

    doc.build(story, onFirstPage=rodape_num, onLaterPages=rodape_num)
    print("PDF gerado:", caminho)

def html_bold(txt):
    return re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', txt)

# ---------------------------------------------------------------------------
# 7. PAGINAÇÃO VIA pypdf
# ---------------------------------------------------------------------------
def pagina_secoes(pdf_path, capitulos):
    from pypdf import PdfReader
    leitor = PdfReader(pdf_path)
    textos = [(i + 1, (p.extract_text() or "")) for i, p in enumerate(leitor.pages)]

    inicio_corpo = None
    for pag, txt in textos:
        if "Vivemos em um mundo" in txt:
            inicio_corpo = pag
            break
    assert inicio_corpo, "página inicial do corpo não encontrada no PDF"

    res = {"introducao": 1}
    for n, *_ in capitulos:
        for pag, txt in textos:
            if pag < inicio_corpo:
                continue
            if re.search(rf'CAPÍTULO {n}(?!\d)', txt):
                res[f"cap{n}"] = pag - inicio_corpo + 1
                break
    return inicio_corpo, res

# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    print("== 1. Conteúdo (leitura/limpeza/estrutura) ==")
    conteudo = montar_conteudo()
    total_pal = 0
    intro_pal = sum(len(p.split()) for _, p in conteudo["introducao"])
    print(f"Introdução: {intro_pal} palavras")
    for i, blocos in enumerate(conteudo["capitulos"], 1):
        pal = sum(len(p.split()) for t, p in blocos if t in ("p", "h2"))
        total_pal += pal
        print(f"Capítulo {i}: {pal} palavras")
    print(f"Total do corpo: {intro_pal + total_pal} palavras")

    print("== 2. PDF passada 1 (sem números, para medir paginação) ==")
    gerar_pdf(conteudo, OUT_PDF, numero_paginas=False)
    inicio, paginas = pagina_secoes(OUT_PDF, CAPITULOS)
    print("Páginas reais:", inicio, paginas)

    print("== 3. PDF final com números ==")
    conteudo["paginas"] = paginas
    gerar_pdf(conteudo, OUT_PDF, numero_paginas=True, inicio_corpo=inicio)

    print("== 4. DOCX final ==")
    gerar_docx(conteudo, OUT_DOCX)

    print("== 5. Verificação cruzada ==")
    # reabre o PDF e confere se os números do sumário batem com as páginas reais
    from pypdf import PdfReader
    _, paginas2 = pagina_secoes(OUT_PDF, CAPITULOS)
    ok = all(paginas[k] == paginas2[k] for k in paginas)
    print("Sumário vs páginas reais:", "OK" if ok else f"DIVERGENTE {paginas} vs {paginas2}")
    print("Páginas totais do PDF:", len(PdfReader(OUT_PDF).pages))
