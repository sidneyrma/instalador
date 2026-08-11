# -*- coding: utf-8 -*-
"""
Normaliza "Evolução da Alma — Caminhos para o Autoconhecimento, Fé e Transformação
Pessoal (27 Correções)" a partir do DOCX.
- Lê parágrafos com estilos (Heading 1 = seções/capítulos, Heading 2 = subseções)
- Remove o "O"/"0" solto da abertura (ornamento antes do Prefácio)
- Ignora o sumário/índice (não é conteúdo do livro)
- Gera JSON estruturado para a página de leitura
"""
import re, json
from pathlib import Path

HERE = Path(__file__).parent
SRC = HERE.parent.parent / "livro" / "Evolução da Alma Caminhos para o Autoconhecimento, Fé e Transformação Pessoal 27 Correções.docx"
OUT_JSON = HERE / "evolucao_v2_dados.json"

import docx

def limpar(t):
    t = t.replace('\u00a0', ' ')
    t = re.sub(r'\s+', ' ', t)
    return t.strip()

def main():
    doc = docx.Document(SRC)
    print(f"DOCX: {len(doc.paragraphs)} parágrafos")

    estrutura = []  # (tipo, texto) — h1_parte, h1, h2, p
    pulando_sumar = False
    par_atual = []

    def flush_par():
        nonlocal par_atual
        if par_atual:
            texto = limpar(' '.join(par_atual))
            if texto:
                estrutura.append(("p", texto))
            par_atual = []

    for i, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        estilo = p.style.name if p.style else "Normal"

        # Remove o "0"/"O" solto da abertura (antes do Prefácio)
        if i == 0 and texto in ("0", "O", "o"):
            print(f"  ↳ Removido ornamento solto: '{texto}' (parágrafo 0)")
            continue

        # Pula a capa/linhas vazias iniciais até o PREFÁCIO
        if i < 3 and not texto:
            continue

        # DETECTA SUMÁRIO: pula até aparecer o primeiro Heading real
        if texto.upper() == "SUMÁRIO":
            pulando_sumar = True
            continue
        if pulando_sumar:
            # sai do sumário quando achar "PREFÁCIO" ou "ABERTURA" ou heading H1
            if estilo.startswith("Heading") and texto.upper() not in ("SUMÁRIO",):
                pulando_sumar = False
                # cai no fluxo normal abaixo
            else:
                continue

        if estilo.startswith("Heading"):
            flush_par()
            nivel = estilo.replace("Heading ", "").strip()
            if nivel == "1":
                # Título de seção/capítulo
                if texto.upper() in ("INÍCIO DO DESENVOLVIMENTO",):
                    estrutura.append(("h1_parte", texto.upper()))
                elif texto.upper() in ("PREFÁCIO", "ABERTURA", "ORAÇÃO FINAL", "FRASE FINAL"):
                    estrutura.append(("h1_parte", texto.upper()))
                elif re.match(r'^Cap[ií]tulo\s*\d+', texto, re.I):
                    estrutura.append(("h1", texto))
                else:
                    # seção introdutória
                    estrutura.append(("h1", texto))
            elif nivel == "2":
                estrutura.append(("h2", texto))
            else:
                estrutura.append(("h2", texto))
        elif texto:
            par_atual.append(texto)
        else:
            flush_par()
    flush_par()

    # Salvar JSON
    dados = {
        "titulo": "Evolução da Alma",
        "subtitulo": "Caminhos para o Autoconhecimento, Fé e Transformação Pessoal",
        "estrutura": [{"tipo": t, "texto": x} for t, x in estrutura],
    }
    OUT_JSON.write_text(json.dumps(dados, ensure_ascii=False, indent=2), encoding='utf-8')

    # Estatísticas
    total_pal = sum(len(x.split()) for t, x in estrutura if t == "p")
    n_h1 = sum(1 for t, _ in estrutura if t == "h1")
    n_h2 = sum(1 for t, _ in estrutura if t == "h2")
    n_parte = sum(1 for t, _ in estrutura if t == "h1_parte")
    print(f"Blocos: {len(estrutura)} | Palavras: {total_pal} | H1: {n_h1} | H2: {n_h2} | Partes: {n_parte}")
    print(f"JSON: {OUT_JSON}")

    # Primeiros blocos (verificar remoção do '0')
    print("\n=== PRIMEIROS 8 BLOCOS ===")
    for t, x in estrutura[:8]:
        print(f"  [{t}] {x[:70]}")

    print("\n=== TÍTULOS H1 (seções e capítulos) ===")
    for t, x in estrutura:
        if t in ("h1", "h1_parte"):
            print(f"  {x[:70]}")

if __name__ == "__main__":
    main()
